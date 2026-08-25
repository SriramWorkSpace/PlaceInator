"""Regression coverage for the shared upload-size guard
(placeinator.api.uploads.read_upload).

/api/resumes, /api/resumes/extract, and /api/jobs/extract previously read an
uploaded file in full via `await file.read()` with no cap -- an oversized or
malformed file could exhaust memory or trigger pathological parse time in
pdfplumber before any format/content validation ever ran. This pins the fix
(413 once the configured limit is exceeded) and that ordinary, valid uploads
are unaffected.

Only the POST /api/resumes accepted-upload case needs the real embedding
model (create_resume embeds every chunk) and is marked accordingly; the
extract endpoints never embed, and every oversized-upload case is rejected
by read_upload before parsing (and, for POST /api/resumes, after the
onboarding check that endpoint already runs first), so none of those need
the model.
"""

from __future__ import annotations

from io import BytesIO

import pytest
import pytest_asyncio
from docx import Document
from httpx import ASGITransport, AsyncClient

from placeinator.api.uploads import MAX_UPLOAD_BYTES
from placeinator.app import create_app
from placeinator.db.migrate import upgrade_to_head
from placeinator.security import generate_token

pytestmark = [pytest.mark.asyncio]

_OVERSIZED = b"a" * (MAX_UPLOAD_BYTES + 1)


def _valid_jd_docx() -> bytes:
    document = Document()
    document.add_paragraph("Backend Engineer at Acme")
    document.add_paragraph("Looking for someone with Python and FastAPI experience.")
    buf = BytesIO()
    document.save(buf)
    return buf.getvalue()


@pytest.fixture
def _sidecar_env(tmp_path, monkeypatch):
    from placeinator.db.session import reset_engine
    from placeinator.settings import get_settings

    monkeypatch.setenv("PLACEINATOR_DATA_DIR", str(tmp_path))
    get_settings.cache_clear()
    reset_engine()
    upgrade_to_head()
    yield
    reset_engine()
    get_settings.cache_clear()


@pytest_asyncio.fixture
async def client(_sidecar_env):
    # Skips app.router.lifespan_context(app) for the same reason as
    # test_matching_api.py's client fixture: avoids kicking off the real
    # background embedding-model warm-up for tests that don't need it.
    # test_resume_upload_accepts_a_normal_file_and_still_creates_the_resume
    # is the one test in this file that genuinely needs the model, and it
    # reaches it lazily through create_resume -> embed_texts, not warm-up.
    token = generate_token()
    app = create_app()
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://sidecar") as c:
        c.headers["Authorization"] = f"Bearer {token}"
        yield c


async def test_resume_extract_accepts_a_normal_upload(client):
    response = await client.post(
        "/api/resumes/extract",
        data={"source_format": "tex"},
        files={"file": ("jane.tex", b"Jane Doe\nSkills: Python, FastAPI", "text/x-tex")},
    )
    assert response.status_code == 200, response.text


async def test_resume_extract_rejects_an_oversized_upload(client):
    response = await client.post(
        "/api/resumes/extract",
        data={"source_format": "tex"},
        files={"file": ("huge.tex", _OVERSIZED, "text/x-tex")},
    )
    assert response.status_code == 413, response.text


async def test_jobs_extract_accepts_a_normal_upload(client):
    response = await client.post(
        "/api/jobs/extract",
        data={"source_format": "docx"},
        files={
            "file": (
                "jd.docx",
                _valid_jd_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200, response.text
    assert "Python" in response.json()["description"]


async def test_jobs_extract_rejects_an_oversized_upload(client):
    response = await client.post(
        "/api/jobs/extract",
        data={"source_format": "pdf"},
        files={"file": ("huge.pdf", _OVERSIZED, "application/pdf")},
    )
    assert response.status_code == 413, response.text


async def test_resume_upload_rejects_an_oversized_file(client):
    onboarding = await client.put(
        "/api/profile",
        json={"full_name": "Jane Doe", "email": "jane@example.com"},
    )
    assert onboarding.status_code == 200

    response = await client.post(
        "/api/resumes",
        data={"label": "X", "source_format": "tex"},
        files={"file": ("huge.tex", _OVERSIZED, "text/x-tex")},
    )
    assert response.status_code == 413, response.text


@pytest.mark.model
async def test_resume_upload_accepts_a_normal_file_and_still_creates_the_resume(client):
    onboarding = await client.put(
        "/api/profile",
        json={"full_name": "Jane Doe", "email": "jane@example.com"},
    )
    assert onboarding.status_code == 200

    response = await client.post(
        "/api/resumes",
        data={"label": "SDE Resume", "source_format": "tex"},
        files={"file": ("jane.tex", b"Jane Doe\nSkills: Python, FastAPI", "text/x-tex")},
    )
    assert response.status_code == 201, response.text
    assert response.json()["chunk_count"] > 0
